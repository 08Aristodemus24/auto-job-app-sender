import os 

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Length

from argparse import ArgumentParser

from typing import List


def load_files(path: str, is_txt=False) -> tuple:
    """
    loads the either an arbitrary file other than a .txt file or a .txt file
    """

    with open(path, 'rb') as file:
        # read binary data of file. If file is .txt then deocde to normal string
        file_data = file.read() if is_txt is False else file.read().decode('ascii')

        # get base name of file before sending
        file_name = os.path.basename(file.name)

    return file_data, file_name

def load_cover_letter(position: str):
    """
    loads the cover letter based on the position argument which tells 
    whether to load the data analyst cover letter or the ML engineer
    cover letter
    """

    cover_letter = {
        "Data Analyst": lambda: load_files('../documents/DA_cover_letter_2.txt', is_txt=True),
        "Machine Learning Engineer": lambda: load_files('../documents/MLE_cover_letter_2.txt', is_txt=True),
        "Software Engineer": lambda: load_files('../documents/SWE_cover_letter_2.txt', is_txt=True),
        "Data Engineer": lambda: load_files('../documents/DE_cover_letter_2.txt', is_txt=True),
        "Data Scientist": lambda: load_files('../documents/DA_cover_letter_2.txt', is_txt=True),
        "Python Developer": lambda: load_files('../documents/SWE_cover_letter_2.txt', is_txt=True),
    }

    file_data, file_name = cover_letter[position]()

    return file_data, file_name

def emphasize(feature: str):
    """
    Capitalizes each letter in a phrase. E.g. creative dynamix solutions, 
    inc. is converted to Creative Dynamix Solutions, Inc.
    """
    feature = feature.strip()
    return " ".join([word.capitalize() for word in feature.split(' ')])

def view_sections(sections: List[str]) -> None:
    for section in sections:
        print(f"{section}")

def main(args):
    """
    
    """
    position = emphasize(feature=args.position)
    company_name = emphasize(feature=args.company_name) 

    # load cover letter file based on arguments then format 
    # content which contains place holders for our arguments
    content, _ = load_cover_letter(position)
    formatted_content = content.format(position=position, company_name=company_name)
    
    # test = load_cover_letter(position)
    # print(test)

    # remove trailing spaces and split on consecutive characters \r\n\r\n
    sections = formatted_content.split('\r\n\r\n')
    # view_sections(sections)

    salutation_sect = sections[0] + '\r\n'
    contents = sections[1:-3]   
    conclusion_sect = sections[-3] + '\r\n'
    close_sect = sections[-2] + '\r\n'
    ps_sect = sections[-1]
    

    
    # instantiate Document() object
    document = Document()
    
    # following lines will change font and size of content
    style = document.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # align left salutation
    salutation = document.add_paragraph(salutation_sect)

    # justify content
    for content in contents:
        content = document.add_paragraph(content)
        
        # justify paragraph which is represented by 3
        content.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # justify conclusion also
    conclusion = document.add_paragraph(conclusion_sect)
    conclusion.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # align left close
    close = document.add_paragraph(close_sect)

    # align left P.S.
    ps = document.add_paragraph(ps_sect)
    print(document)
    
    pos_placeholder = position.replace(' ', '-')
    comp_placeholder = company_name.replace(' ', '-') 

    # save document
    document.save(f'../documents/cover-letter_{pos_placeholder}_{comp_placeholder}.docx')

